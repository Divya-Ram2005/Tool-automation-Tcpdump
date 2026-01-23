import getpass
import yaml
import os
import subprocess
from datetime import datetime
from docx import Document
import re
import paramiko
import shlex
import time

# auto-detect a network interface that is UP and has a non-loopback IP
def autodetect_interface(remote_ssh=None):
    # returns interface name, e.g. "ens33" or None
    cmd = "ip -o -4 addr show scope global | awk '{print $2, $4}'"
    out, err = run_command(cmd, "", remote_ssh)
    if not out:
        return None
    # parse output lines like: "ens33 192.168.1.20/24"
    for line in out.splitlines():
        parts = line.strip().split()
        if len(parts) >= 2:
            iface = parts[0]
            ip = parts[1].split("/")[0]
            if ip != "127.0.0.1":
                return iface
    return None

# run a list of commands (pre_actions), replacing {iface} and {timeout}
def run_pre_actions(actions, iface, timeout, sudo_password, remote_ssh=None):
    if not actions:
        return
    for c in actions:
        cmd = c.replace("{iface}", iface or "").replace("{timeout}", str(timeout))
        print(f"[PRE] {cmd}")
        out, err = run_command(cmd, sudo_password, remote_ssh)
        print(out, err)

# safe wrapper to format check_command with iface and timeout
def format_cmd(template, iface, timeout):
    cmd = template.replace("{iface}", iface or "").replace("{timeout}", str(timeout))
    return cmd


TCPDUMP_LATEST = "4.99.5"
LIBPCAP_LATEST = "1.10.5"

# ---------------- Helper: Compare Versions ----------------
def parse_version(v):
    """Return list of ints for version string like '4.99.5' or None if not parseable."""
    if not v or not isinstance(v, str):
        return None
    m = re.search(r"\d+(?:\.\d+)+", v)
    if not m:
        return None
    parts = m.group(0).split(".")
    try:
        return [int(x) for x in parts]
    except ValueError:
        return None

def compare_versions(v1, v2):
    """
    Compare two version strings.
    Return -1 if v1 < v2, 0 if equal, 1 if v1 > v2.
    If unparsable, treat version as older.
    """
    p1 = parse_version(v1)
    p2 = parse_version(v2)

    if p1 is None and p2 is None:
        return 0
    if p1 is None:
        return -1
    if p2 is None:
        return 1

    # pad shorter version
    maxlen = max(len(p1), len(p2))
    p1 += [0] * (maxlen - len(p1))
    p2 += [0] * (maxlen - len(p2))

    if p1 < p2:
        return -1
    if p1 > p2:
        return 1
    return 0
# ---------------- Helper to run commands ----------------
def run_command(cmd, sudo_password="", remote_ssh=None):
    if remote_ssh:
        if "sudo " in cmd:
            cmd = cmd.replace("sudo ", f"echo '{sudo_password}' | sudo -S ")
        stdin, stdout, stderr = remote_ssh.exec_command(cmd, get_pty=True)
        out_lines, err_lines = [], []
        for line in iter(stdout.readline, ""):
            out_lines.append(line)
        for line in iter(stderr.readline, ""):
            err_lines.append(line)
        return "".join(out_lines), "".join(err_lines)
    else:
        if "sudo " in cmd:
            cmd = cmd.replace("sudo ", f"echo '{sudo_password}' | sudo -S ")
        proc = subprocess.run(cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        return proc.stdout, proc.stderr

# ---------------- TCPDump Version Check ----------------
def get_tcpdump_version(local=True, ssh=None):
    """
    Return the tcpdump numeric version string like '4.99.5', or None if not found.
    Uses regex to extract version.
    """
    cmd = "tcpdump --version | head -n 1"

    if ssh:
        out, err = run_command(cmd, remote_ssh=ssh)
    else:
        out, err = run_command(cmd)

    combined = (out or "") + "\n" + (err or "")
    match = re.search(r"\d+(?:\.\d+)+", combined)
    if match:
        return match.group(0)
    return None
# ---------------- Installation Logic ----------------
def install_tcpdump(sudo_password, ssh=None):
    print("[INFO] Installing tcpdump...")

    apt_cmds = [
        "sudo apt-get update",
        "sudo apt-get install -y tcpdump libpcap-dev build-essential flex bison wget"
    ]

    for c in apt_cmds:
        out, err = run_command(c, sudo_password, ssh)
        if "error" in err.lower():
            break
    else:
        return True

    # Fallback: Source Installation
    print("[WARN] apt install failed, installing via source")

    libpcap_url = f"https://www.tcpdump.org/release/libpcap-{LIBPCAP_LATEST}.tar.gz"
    tcpdump_url = f"https://www.tcpdump.org/release/tcpdump-{TCPDUMP_LATEST}.tar.gz"

    cmds = [
        f"wget -q {libpcap_url} -O /tmp/libpcap.tar.gz",
        "tar -xzf /tmp/libpcap.tar.gz -C /tmp",
        f"cd /tmp/libpcap-{LIBPCAP_LATEST} && ./configure && make && sudo make install",
        f"wget -q {tcpdump_url} -O /tmp/tcpdump.tar.gz",
        "tar -xzf /tmp/tcpdump.tar.gz -C /tmp",
        f"cd /tmp/tcpdump-{TCPDUMP_LATEST} && ./configure && make && sudo make install"
    ]

    for c in cmds:
        out, err = run_command(c, sudo_password, ssh)

    return True

# ---------------- MAIN ----------------
print("\nChoose execution mode:")
print("1) Local Machine")
print("2) Remote Machine")
choice = input("Enter choice [1-2]: ").strip()

use_remote = (choice == "2")

ssh = None
if use_remote:
    ip = input("Remote IP: ").strip()
    user = input("Username: ").strip()
    ssh_pass = getpass.getpass("SSH Password: ")
    sudo_pass = getpass.getpass("Remote sudo password: ")
    engineer = input("Engineer Name: ").strip() or "Engineer"

    ssh = paramiko.SSHClient()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(ip, username=user, password=ssh_pass)
else:
    engineer = input("Engineer Name: ").strip() or "Engineer"
    sudo_pass = getpass.getpass("Local sudo password: ")

yaml_file = "testcases_tcpdump.yaml"
testcases = yaml.safe_load(open(yaml_file, "r"))['testcases']

results = []

# Check TCPDump version
current_version = get_tcpdump_version(local=not use_remote, ssh=ssh)
if not current_version or compare_versions(current_version, TCPDUMP_LATEST) < 0:
    print("[WARN] tcpdump not installed OR outdated → installing")
    install_tcpdump(sudo_pass, ssh)
    current_version = get_tcpdump_version(local=not use_remote, ssh=ssh)

# ---------------- Run Testcases ----------------
# detect interface once
detected_iface = None
if use_remote:
    detected_iface = autodetect_interface(ssh)
else:
    detected_iface = autodetect_interface(None)

print(f"[INFO] Using interface: {detected_iface or 'auto/unspecified'}")

for tc in testcases:
    iface = tc.get("interface") or detected_iface
    timeout = tc.get("timeout", 8)

    print(f"\n[INFO] Running {tc['id']}: {tc['description']} (iface={iface}, timeout={timeout})")


    # format check_command (tcpdump)
    check_cmd = format_cmd(tc["check_command"], iface or "", timeout)
 
    # Start tcpdump first in background
    print(f"[INFO] Starting tcpdump: {check_cmd}")
    tcpdump_proc = subprocess.Popen(check_cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

    # small sleep to ensure tcpdump is running
    time.sleep(1)

    # Run pre-actions to generate traffic (if any)
    run_pre_actions(tc.get("pre_actions", []), iface, timeout, sudo_pass, ssh if use_remote else None)

    out_lines = []
    err_lines = []

    # Read stdout live (tcpdump packets)
    for line in iter(tcpdump_proc.stdout.readline, ""):
        print(line, end="")          # print to terminal
        out_lines.append(line)       # store for DOCX

    # Read stderr (tcpdump summary, warnings)
    for line in iter(tcpdump_proc.stderr.readline, ""):
        print(line, end="")
        err_lines.append(line)

    tcpdump_proc.wait()

    out = "".join(out_lines)
    err = "".join(err_lines)
    
    passed = tc["expected_output_substring"] in out
    
    # Auto-generate notes depending on output
    if passed:
        note = "Test passed. Expected pattern detected in output."
    else:
        if out.strip() == "":
            note = "Test failed. No packet data captured or timeout occurred."
        else:
            note = "Test failed. Output present but expected pattern not detected."
    #store results
    results.append({
        "id": tc["id"],
        "description": tc["description"],
        "output": out.strip(),
        "error": err.strip(),
        "passed": passed,
        "notes": note

    })

# ---------------- Generate DOCX Report ----------------
from docx import Document
from datetime import datetime

doc = Document()
doc.add_heading("TCPDump Automated Test Report", level=0)

doc.add_paragraph(f"Engineer: {engineer}")
doc.add_paragraph(f"Date: {datetime.now().strftime('%d.%m.%Y')}")
doc.add_paragraph(f"TCPDump Version: {current_version}")
doc.add_paragraph(" ")

# Add a summary table
doc.add_heading("Test Summary", level=1)
summary_table = doc.add_table(rows=1, cols=4)
hdr = summary_table.rows[0].cells
hdr[0].text = "Test ID"
hdr[1].text = "Description"
hdr[2].text = "Status"
hdr[3].text = "Expected String"

for tc, r in zip(testcases, results):
    row = summary_table.add_row().cells
    row[0].text = r["id"]
    row[1].text = r["description"]
    row[2].text = "PASS" if r["passed"] else "FAIL"
    row[3].text = tc["expected_output_substring"]

doc.add_page_break()

# Detailed test sections
doc.add_heading("Detailed Test Results", level=1)

for tc, r in zip(testcases, results):

    doc.add_heading(r["id"], level=2)
    doc.add_paragraph(f"Description: {tc['description']}")

    # Pre-actions
    pre = tc.get("pre_actions", [])
    if pre:
        doc.add_paragraph("Pre-Actions:", style="List Bullet")
        for p in pre:
            doc.add_paragraph(p, style="List Continue")
    else:
        doc.add_paragraph("Pre-Actions: None")

    # Check command
    doc.add_paragraph(f"Check Command:\n{tc['check_command']}")

    # Expected output
    doc.add_paragraph(f"Expected Output Substring: {tc['expected_output_substring']}")

    # Actual result
    doc.add_paragraph("Actual Output:")
    doc.add_paragraph(r["output"] or "NO OUTPUT")

    # Errors
    doc.add_paragraph("Error Output:")
    doc.add_paragraph(r["error"] or "NO ERROR")

    # Status
    doc.add_paragraph(f"Final Status: {'PASS' if r['passed'] else 'FAIL'}")

    # Notes (NEW SECTION)
    notes = tc.get("notes", None)
    if notes:
        doc.add_paragraph("Notes:")
        doc.add_paragraph(notes)
    else:
        doc.add_paragraph("Notes: None")
    
    doc.add_paragraph("----------------------------------------")

# Save file
file = f"TCPDump_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
doc.save(file)

print(f"\n[INFO] Report saved as {file}")
